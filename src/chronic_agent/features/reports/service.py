from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session

from chronic_agent.core.contracts import ReportOut, ReportRequest
from chronic_agent.features.clinician_digest.service import ClinicianDigestService
from chronic_agent.features.companion.service import CompanionService
from chronic_agent.features.meals.service import MealService
from chronic_agent.platform.repositories import ReportRepository
from chronic_agent.platform.settings import settings


class ReportService:
    def __init__(self, db: Session, patient_id: int):
        self.repo = ReportRepository(db, patient_id)
        self.companion = CompanionService(db, patient_id)
        self.meals = MealService(db, patient_id)
        self.digest = ClinicianDigestService(db, patient_id)
        self.export_dir = settings.export_path

    def _compose(self, payload: ReportRequest) -> tuple[str, str]:
        if payload.report_type == 'patient_weekly':
            today = self.companion.today_view()
            meal = self.meals.weekly_summary()
            title = '患者周报'
            markdown = '\n'.join([
                '# 患者周报',
                '',
                f"- 最新空腹血糖：{today.latest_fasting_glucose or '暂无'}",
                f"- 最新血压：{today.latest_blood_pressure or '暂无'}",
                f'- 待处理提醒：{today.pending_reminders}',
                f"- 本周饮食高风险标签：{', '.join(meal.top_risk_tags) if meal.top_risk_tags else '暂无明显高风险'}",
                '',
                f'- 陪伴建议：{today.coach_message}',
            ])
        elif payload.report_type == 'adherence_overview':
            today = self.companion.today_view()
            title = '依从性概览'
            markdown = '\n'.join([
                '# 依从性概览',
                '',
                f'- 今日待处理提醒：{today.pending_reminders}',
                '- 建议重点查看是否存在连续跳过服药或长期未记录监测值的情况。',
            ])
        else:
            digest = self.digest.generate(payload.window_days)
            title = '门诊前摘要报告'
            markdown = digest.content_markdown
        return title, markdown

    def generate(self, payload: ReportRequest) -> ReportOut:
        title, markdown = self._compose(payload)
        row = self.repo.add(payload.report_type, payload.window_days, title, markdown)
        return ReportOut(id=row.id, patient_id=row.patient_id, report_type=row.report_type, window_days=row.window_days, title=row.title, markdown=row.markdown, created_at=row.created_at)

    def list_reports(self) -> list[ReportOut]:
        rows = self.repo.list_recent()
        return [ReportOut(id=r.id, patient_id=r.patient_id, report_type=r.report_type, window_days=r.window_days, title=r.title, markdown=r.markdown, created_at=r.created_at) for r in rows]

    def export(self, report_id: int, fmt: str) -> Path:
        row = self.repo.get(report_id)
        if row is None:
            raise ValueError('report not found')
        safe = f'report_{report_id}_{fmt}'
        if fmt == 'html':
            path = self.export_dir / f'{safe}.html'
            html = '<html><body><pre style="white-space: pre-wrap; font-family: Arial;">' + row.markdown + '</pre></body></html>'
            path.write_text(html, encoding='utf-8')
            self.repo.set_export_paths(report_id, html_path=str(path))
            return path
        if fmt == 'docx':
            path = self.export_dir / f'{safe}.docx'
            doc = Document()
            for line in row.markdown.splitlines():
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                else:
                    doc.add_paragraph(line)
            doc.save(path)
            self.repo.set_export_paths(report_id, docx_path=str(path))
            return path
        if fmt == 'pdf':
            path = self.export_dir / f'{safe}.pdf'
            c = canvas.Canvas(str(path), pagesize=A4)
            width, height = A4
            y = height - 40
            for line in row.markdown.splitlines():
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:95])
                y -= 16
            c.save()
            self.repo.set_export_paths(report_id, pdf_path=str(path))
            return path
        raise ValueError('unsupported export format')
